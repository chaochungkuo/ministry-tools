from docx import Document
from docx.shared import RGBColor, Inches
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.enum.text import WD_COLOR_INDEX

def insert_paragraph_after(paragraph, text, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return run, new_para

class TextBody():
    def __init__(self, path, format="book"):
        """type = outline | book
        """
        self.path = path
        self.ext = path.split(".")[-1]
        self.format = format
        self.doc = None
        self.texts = []
        self.scriptures = []
        self.outline_label = []
        self.scriptures_str = []
        self.load()

    def load(self):
        if self.ext == "docx":
            self.doc = Document(self.path)
            for para in self.doc.paragraphs:
                if self.format == "book":
                    self.texts.append(para.text.strip())
                # if self.format == "outline":
                #     l = para.text.strip().split()
                #     self.outline_label.append(l[0].replace(".", ""))
                #     self.texts.append(l[1])
        elif self.ext == "txt":
            with open(self.path) as f:
                for line in f:
                    if self.format == "book":
                        self.texts.append(line.strip())
                    # if self.format == "outline":
                    #     l = line.strip().split()
                    #     self.outline_label.append(l[0].replace(".", ""))
                    #     self.texts.append(l[1])

    def save_docx(self, path, ref=False, de=False):
        # document = Document()
        if not self.scriptures_str:
            print("*** Please prepare scripture outputs.")
        
        if de and not self.de_lines:
            print("*** Please prepare de lines.")

        for i, p in enumerate(self.doc.paragraphs):
            # document.add_paragraph(p.text, p.style)
            # print(len(document.paragraphs))
            if de:
                r, de_p = insert_paragraph_after(p, self.de_lines[i], style=p.style)
                font = r.font
                font.highlight_color = WD_COLOR_INDEX.YELLOW
            if ref:
                if de:
                    r, newp = insert_paragraph_after(de_p, self.scriptures_str[i])
                    
                    

                else:
                    r, newp = insert_paragraph_after(p, self.scriptures_str[i])
                    # Set the paragraph formatting
                    # paragraph_format = newp.paragraph_format
                    # paragraph_format.first_line_indent = Inches(-0.5)  # First line outdent
                    # paragraph_format.left_indent = Inches(0.5)  # Indentation for the rest of the paragraph
                    # Add tab stop to align the text after the reference
                    # tab_stops = newp.paragraph_format.tab_stops
                    # tab_stops.add_tab_stop(Inches(3.0))  # Adjust as per your preferred tab alignment
                # print(r)
                # r = document.add_paragraph().add_run(self.scriptures[i])
                font = r.font
                font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
                # newp.paragraph_format.left_indent = Inches(0.5)
                # newp.paragraph_format.first_line_indent = Inches(-0.5)
        self.doc.save(path)
